"""
유튜브 쇼츠/인스타 릴스 자동 생성 모듈 — '오늘의 마감임박 TOP 5'
────────────────────────────────────────────────────────────
크롤링 완료 직후 하루 1회 실행:
  ① 마감 임박 공고 TOP 5 선별 (공고종료일 오름차순, 기관 중복 제거)
  ② MZ 톤 스크립트 자동 생성 (매일 문구 로테이션)
  ③ Google Cloud TTS로 생동감 있는 나레이션 합성 (실패 시 무음으로 진행)
  ④ 장면 렌더링: 인트로 → 공고별 5장면(기관 CI/기관명/제목/D-day) → 아웃트로
  ⑤ ffmpeg(imageio-ffmpeg 정적 바이너리)로 1080×1920 세로 영상 합성
     - 장면마다 줌 모션(Ken Burns), 나레이션 + 배경음악(BGM) 믹싱
  ⑥ 결과물 업로드:
     - Google Drive '[shorts]' 폴더 (컴퓨터가 꺼져 있어도 클라우드에 저장,
       PC를 켜면 'G:\\내 드라이브\\...' 경로로 자동 동기화됨)
     - 인스타그램 릴스 (기존 _config/social 토큰 재사용)

■ 설정 위치: Firestore `_config` 컬렉션 → `shorts` 문서
    enabled          : false 로 두면 전체 비활성 (기본: 문서 없으면 비활성)
    drive_folder_id  : 드라이브 '[shorts]' 폴더 ID (폴더 URL의 마지막 경로 조각)
                       ※ 이 폴더를 함수 서비스계정 이메일에 '편집자'로 공유해야 함
    ig_reels         : true 면 인스타그램 릴스도 게시 (기본 true)
    tts_voice        : (선택) 기본 'ko-KR-Neural2-A'
    last_run         : (자동 관리) 마지막 실행 날짜 — 하루 1회 보장

■ BGM: assets/bgm.mp3 가 있으면 그 파일을 사용(유튜브 오디오 라이브러리 등
  무료 음원을 넣어두면 됨). 없으면 저작권 걱정 없는 자체 합성 루프를 생성해 사용.
"""

import io
import math
import os
import random
import re
import struct
import subprocess
import tempfile
import time
import wave
from datetime import datetime, timedelta, timezone

import requests

KST = timezone(timedelta(hours=9))
SITE_URL = "https://www.allgongin.com"
FONT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "fonts")
ASSET_BGM = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets", "bgm.mp3")

W, H = 1080, 1920          # 쇼츠 규격 (9:16)
FPS = 30
SCENE_MIN = 3.2            # 장면 최소 길이(초)
SCENE_MAX = 10.0           # 장면 최대 길이(초) — 나레이션이 끝까지 재생되도록 여유 확보
TTS_RATE = 1.16            # 기본 말 속도
TTS_RATE_MAX = 1.45        # 문장이 길어 장면 한도를 넘을 때 올릴 수 있는 최대 속도
INTRO_MIN, OUTRO_MIN = 2.2, 2.6

# 순위별 포인트 컬러 (장면마다 화면 분위기가 바뀌도록)
RANK_COLORS = ["#e8590c", "#1c7ed6", "#2b8a3e", "#9c36b5", "#e03131"]


def _today_kst() -> str:
    return datetime.now(KST).strftime("%Y-%m-%d")


def _ffmpeg() -> str:
    import imageio_ffmpeg
    return imageio_ffmpeg.get_ffmpeg_exe()


# ─────────────────────── ① 마감임박 TOP 5 선별 ───────────────────────
def select_top5(df):
    """공고종료일이 오늘 이후인 공고를 마감 가까운 순으로, 기관 중복 없이 5건."""
    import pandas as pd
    today = pd.Timestamp(datetime.now(KST).date())
    d = df[df["공고종료일"].notna() & (df["공고종료일"] >= today)].sort_values("공고종료일")
    seen, picked = set(), []
    for _, r in d.iterrows():
        org = str(r.get("기관명", "")).strip()
        if not org or org in seen:
            continue
        seen.add(org)
        picked.append(r)
        if len(picked) == 5:
            break
    return picked


def _dday(end) -> str:
    try:
        days = (end.date() - datetime.now(KST).date()).days
        if days <= 0:
            return "D-DAY"
        return f"D-{days}"
    except Exception:  # noqa: BLE001
        return ""


def _dday_spoken(end) -> str:
    """나레이션용 D-day 문구 (템플릿의 어느 자리에 넣어도 자연스럽게 끝나는 형태)."""
    try:
        days = (end.date() - datetime.now(KST).date()).days
        if days <= 0:
            return "마감이 바로 오늘"
        if days == 1:
            return "마감이 바로 내일"
        return f"마감까지 딱 {days}일"
    except Exception:  # noqa: BLE001
        return "마감 임박"


# ─────────────────────── ② MZ 톤 스크립트 생성 ───────────────────────
INTRO_HOOKS = [
    "스크롤 멈춰! 오늘 놓치면 사라지는 공공기관 채용 TOP 파이브!",
    "잠깐만, 이거 안 보면 손해! 마감 임박 공기업 채용 다섯 개 바로 갑니다!",
    "취준생 모여라! 오늘 마감 임박 꿀공고 TOP 파이브, 삼십 초 요약!",
    "이 영상 뜬 순간이 지원 타이밍! 마감 임박 채용 TOP 파이브!",
    "저장 필수! 지금 아니면 못 쓰는 공공기관 공고 다섯 개 정리해드림!",
]
# 형식: "{N}위는 {기관명}입니다. {공고내용}입니다. {독려 멘트}"
JOB_TEMPLATE = "{rank}위는 {org}입니다. {title}입니다. {cta}"
# 지원 독려 멘트 풀 — 한 영상 안에서는 서로 다른 멘트가 뽑히고, 날마다 조합이 바뀜
JOB_CTAS = [
    "얼른 지원해야겠죠?",
    "지원을 깜빡하면 안 되겠죠?",
    "이건 못 참죠, 바로 지원 각!",
    "놓치면 두고두고 생각날걸요?",
    "마감 전에 클릭, 잊지 마세요!",
    "기회는 준비된 사람이 잡는 법이죠!",
    "오늘의 할 일, 지원서 제출!",
    "고민하는 사이에 마감됩니다!",
    "일단 지원! 후회는 없습니다!",
    "합격의 주인공, 바로 당신일지도요?",
]
OUTRO_LINES = [
    "상세 공고는 올공인닷컴에서! 좋아요 누르고 최종합격 가즈아!",
    "더 많은 공고는 올공인닷컴! 팔로우하면 매일 꿀공고 떠요!",
    "지원 방법은 올공인닷컴에서 확인! 저장해두고 두고두고 보세요!",
]


def _clean_for_speech(text: str, limit: int = 38) -> str:
    """나레이션용 제목 다듬기: 괄호/특수문자 정리 + 길이 제한."""
    s = re.sub(r"[\[\](){}<>]", " ", str(text or ""))
    s = re.sub(r"\s+", " ", s).strip()
    if len(s) > limit:
        s = s[:limit].rstrip() + ""
    return s


def build_script(rows) -> dict:
    """{'intro': str, 'jobs': [str×5], 'outro': str} — 날짜 시드로 매일 로테이션."""
    rnd = random.Random(_today_kst())  # 같은 날은 같은 문구(재실행 대비), 날마다 변화
    ctas = rnd.sample(JOB_CTAS, min(len(rows), len(JOB_CTAS)))  # 영상 안에서 멘트 중복 없음
    jobs = []
    for i, r in enumerate(rows):
        jobs.append(JOB_TEMPLATE.format(
            rank=i + 1,  # TTS가 '1위'를 '일위'로 자연스럽게 읽음
            org=str(r.get("기관명", "")),
            title=_clean_for_speech(r.get("채용공고제목", "")),
            cta=ctas[i % len(ctas)],
        ))
    return {
        "intro": rnd.choice(INTRO_HOOKS),
        "jobs": jobs,
        "outro": rnd.choice(OUTRO_LINES),
    }


# ─────────────────────── ③ TTS 나레이션 (실패해도 영상은 계속) ───────────────────────
DEFAULT_VOICE = "ko-KR-Chirp3-HD-Leda"   # 구글 최신 고품질 음성 (자연스러운 대화체)
FALLBACK_VOICE = "ko-KR-Neural2-A"       # 위 음성 사용 불가 시 대체


def _tts_once(text: str, out_path: str, voice_name: str, rate: float) -> bool:
    from google.cloud import texttospeech
    client = texttospeech.TextToSpeechClient()
    audio_kwargs = {"audio_encoding": texttospeech.AudioEncoding.MP3, "speaking_rate": rate}
    if "Chirp" not in voice_name:
        audio_kwargs["pitch"] = 1.5  # Chirp 계열은 pitch 미지원 → 기존 음성에만 적용
    resp = client.synthesize_speech(
        input=texttospeech.SynthesisInput(text=text),
        voice=texttospeech.VoiceSelectionParams(language_code="ko-KR", name=voice_name),
        audio_config=texttospeech.AudioConfig(**audio_kwargs),
    )
    with open(out_path, "wb") as f:
        f.write(resp.audio_content)
    return True


def synth_tts(text: str, out_path: str, voice_name: str, rate: float = TTS_RATE) -> bool:
    """Google Cloud TTS → mp3 파일. 지정 음성 실패 시 대체 음성으로 재시도."""
    try:
        return _tts_once(text, out_path, voice_name, rate)
    except Exception as e:  # noqa: BLE001
        print(f"  (음성 '{voice_name}' 실패 → 대체 음성 시도: {e})")
    try:
        return _tts_once(text, out_path, FALLBACK_VOICE, rate)
    except Exception as e:  # noqa: BLE001
        print(f"  (TTS 실패 → 해당 장면 무음 진행: {e})")
        return False


def _audio_duration(path: str) -> float:
    """ffmpeg로 오디오 길이(초) 측정 (디코딩 진행시간 → 실패 시 Duration 헤더)."""
    try:
        p = subprocess.run([_ffmpeg(), "-i", path, "-f", "null", "-"],
                           capture_output=True, text=True, timeout=60)
        m = re.findall(r"time=(\d+):(\d+):(\d+\.?\d*)", p.stderr)
        if m:
            h, mnt, s = m[-1]
            return int(h) * 3600 + int(mnt) * 60 + float(s)
        m = re.search(r"Duration:\s*(\d+):(\d+):(\d+\.?\d*)", p.stderr)
        if m:
            return int(m.group(1)) * 3600 + int(m.group(2)) * 60 + float(m.group(3))
    except Exception:  # noqa: BLE001
        pass
    return 0.0


# ─────────────────────── ④ 장면 렌더링 (PIL) ───────────────────────
def _fonts():
    from PIL import ImageFont
    xb = lambda s: ImageFont.truetype(os.path.join(FONT_DIR, "NanumGothicExtraBold.ttf"), s)  # noqa: E731
    rg = lambda s: ImageFont.truetype(os.path.join(FONT_DIR, "NanumGothic.ttf"), s)  # noqa: E731
    return xb, rg


def _grad(img, top_rgb, bottom_rgb):
    """세로 그라데이션 배경."""
    from PIL import ImageDraw
    d = ImageDraw.Draw(img)
    for y in range(H):
        t = y / H
        c = tuple(int(top_rgb[i] + (bottom_rgb[i] - top_rgb[i]) * t) for i in range(3))
        d.line([(0, y), (W, y)], fill=c)


def _hex(c):
    c = c.lstrip("#")
    return tuple(int(c[i:i + 2], 16) for i in (0, 2, 4))


def _wrap(d, text, font, max_w, max_lines=3):
    """픽셀 폭 기준 줄바꿈."""
    words = str(text).split()
    lines, cur = [], ""
    for w_ in words:
        t = (cur + " " + w_).strip()
        if d.textlength(t, font=font) <= max_w:
            cur = t
        else:
            if cur:
                lines.append(cur)
            cur = w_
        if len(lines) == max_lines:
            break
    if cur and len(lines) < max_lines:
        lines.append(cur)
    if len(lines) == max_lines and words and d.textlength(lines[-1], font=font) > max_w:
        lines[-1] = lines[-1][:-1] + ""
    return lines


def _fetch_logo(url):
    from PIL import Image
    try:
        if not url or not str(url).startswith("http"):
            return None
        r = requests.get(url, timeout=10)
        r.raise_for_status()
        return Image.open(io.BytesIO(r.content)).convert("RGBA")
    except Exception:  # noqa: BLE001
        return None


def render_intro(path: str):
    from PIL import Image, ImageDraw
    xb, rg = _fonts()
    img = Image.new("RGB", (W, H))
    _grad(img, _hex("#0b1957"), _hex("#1c7ed6"))
    d = ImageDraw.Draw(img)

    def center(t, f, y, fill="#ffffff"):
        d.text(((W - d.textlength(t, font=f)) / 2, y), t, font=f, fill=fill)

    now = datetime.now(KST)
    center(now.strftime("%m월 %d일").replace("월 0", "월 "), rg(52), 560, "#a5d8ff")
    center("오늘 마감임박", xb(120), 700)
    center("채용 TOP 5", xb(150), 860, "#ffd43b")
    d.rounded_rectangle([(W / 2 - 260, 1120), (W / 2 + 260, 1210)], 45, fill="#e03131")
    center("놓치면 끝!", xb(56), 1138)
    center("공공기관 채용포털 올공 ALLGONG", rg(40), 1720, "#a5d8ff")
    img.save(path, "PNG")


def render_job_scene(path: str, row, rank: int):
    from PIL import Image, ImageDraw
    xb, rg = _fonts()
    accent = RANK_COLORS[(rank - 1) % len(RANK_COLORS)]
    img = Image.new("RGB", (W, H), "#f8f9fa")
    d = ImageDraw.Draw(img)
    # 상단 컬러 밴드 + 순위
    d.rectangle([(0, 0), (W, 380)], fill=_hex(accent))
    d.text((60, 80), f"{rank}", font=xb(220), fill="#ffffff")
    d.text((300, 130), "마감임박", font=xb(64), fill="#ffffff")
    d.text((300, 220), "TOP 5", font=xb(64), fill="#ffd43b")
    # 기관 CI 카드
    card_y = 470
    d.rounded_rectangle([(90, card_y), (W - 90, card_y + 430)], 30, fill="#ffffff",
                        outline=_hex("#dee2e6"), width=3)
    logo = _fetch_logo(row.get("imageUrl") or row.get("로고URL"))
    org = str(row.get("기관명", ""))
    if logo:
        logo.thumbnail((W - 300, 300))
        img.paste(logo, (int((W - logo.width) / 2), card_y + int((430 - logo.height) / 2)), logo)
    else:
        # 로고가 없으면 기관명 이니셜 원 + 기관명
        d.ellipse([(W / 2 - 110, card_y + 60), (W / 2 + 110, card_y + 280)], fill=_hex(accent))
        ini = org[:1] or "공"
        f_ini = xb(120)
        d.text(((W - d.textlength(ini, font=f_ini)) / 2, card_y + 100), ini, font=f_ini, fill="#ffffff")
        f_o = xb(56)
        d.text(((W - d.textlength(org, font=f_o)) / 2, card_y + 310), org, font=f_o, fill="#333333")
    # 기관명 + 공고 제목
    y = 1000
    f_org = xb(72 if len(org) <= 12 else 58)
    d.text(((W - d.textlength(org, font=f_org)) / 2, y), org, font=f_org, fill="#212529")
    y += 120
    f_t = rg(52)
    for ln in _wrap(d, row.get("채용공고제목", ""), f_t, W - 200, 3):
        d.text(((W - d.textlength(ln, font=f_t)) / 2, y), ln, font=f_t, fill="#495057")
        y += 74
    # 고용유형 칩 + D-day 뱃지
    emp = str(row.get("고용유형", "") or "").split(",")[0].strip()
    if emp:
        f_e = rg(44)
        ew = d.textlength(emp, font=f_e) + 70
        d.rounded_rectangle([((W - ew) / 2, 1400), ((W + ew) / 2, 1480)], 40,
                            outline=_hex(accent), width=4)
        d.text(((W - d.textlength(emp, font=f_e)) / 2, 1414), emp, font=f_e, fill=_hex(accent))
    dd = _dday(row.get("공고종료일"))
    f_d = xb(120)
    dw = d.textlength(dd, font=f_d) + 140
    d.rounded_rectangle([((W - dw) / 2, 1540), ((W + dw) / 2, 1730)], 40, fill=_hex("#e03131"))
    d.text(((W - d.textlength(dd, font=f_d)) / 2, 1568), dd, font=f_d, fill="#ffffff")
    f_s = rg(38)
    d.text(((W - d.textlength("allgongin.com", font=f_s)) / 2, 1800), "allgongin.com",
           font=f_s, fill="#adb5bd")
    img.save(path, "PNG")


def render_outro(path: str):
    from PIL import Image, ImageDraw
    xb, rg = _fonts()
    img = Image.new("RGB", (W, H))
    _grad(img, _hex("#1c7ed6"), _hex("#0b1957"))
    d = ImageDraw.Draw(img)

    def center(t, f, y, fill="#ffffff"):
        d.text(((W - d.textlength(t, font=f)) / 2, y), t, font=f, fill=fill)

    center("상세 공고 · 지원 방법", rg(56), 640, "#a5d8ff")
    center("allgongin.com", xb(110), 760, "#ffd43b")
    center("올공 ALLGONG", xb(72), 960)
    d.rounded_rectangle([(W / 2 - 330, 1140), (W / 2 + 330, 1240)], 50, fill="#e03131")
    center("좋아요 · 팔로우 · 저장", xb(52), 1162)
    center("공공기관 채용의 모든 것", rg(44), 1700, "#a5d8ff")
    img.save(path, "PNG")


# ─────────────────────── BGM (무료: 자체 합성, assets/bgm.mp3 있으면 그걸 사용) ───────────────────────
def make_bgm(path_wav: str, seconds: float):
    """저작권 무관한 경쾌한 신스 루프를 직접 합성 (numpy)."""
    import numpy as np
    sr = 44100
    bpm = 122
    beat = 60.0 / bpm
    total = int(sr * seconds)
    t = np.arange(total) / sr
    out = np.zeros(total, dtype=np.float64)
    # 코드 진행: A - E - F#m - D (I-V-vi-IV), 마디마다 순환
    chords = [
        [220.00, 277.18, 329.63],
        [164.81, 207.65, 246.94],
        [185.00, 220.00, 277.18],
        [146.83, 185.00, 220.00],
    ]
    bar = beat * 4
    for i in range(int(seconds / bar) + 1):
        s0, s1 = int(i * bar * sr), min(int((i + 1) * bar * sr), total)
        if s0 >= total:
            break
        seg_t = t[s0:s1] - t[s0]
        pad = np.zeros(s1 - s0)
        for f in chords[i % 4]:
            pad += 0.16 * np.sin(2 * np.pi * f * seg_t) * (1 - 0.3 * np.cos(2 * np.pi * seg_t / bar))
            pad += 0.05 * np.sign(np.sin(2 * np.pi * (f / 2) * seg_t))  # 베이스(스퀘어)
        out[s0:s1] += pad
    # 킥(4비트) + 하이햇(8비트)
    for k in range(int(seconds / beat) + 1):
        s0 = int(k * beat * sr)
        n = min(int(0.12 * sr), total - s0)
        if n <= 0:
            break
        seg = np.arange(n) / sr
        out[s0:s0 + n] += 0.5 * np.sin(2 * np.pi * (120 - 300 * seg) * seg) * np.exp(-seg * 30)
        h0 = int((k + 0.5) * beat * sr)
        hn = min(int(0.05 * sr), total - h0)
        if hn > 0:
            rng = np.random.default_rng(k)
            out[h0:h0 + hn] += 0.08 * rng.standard_normal(hn) * np.exp(-np.arange(hn) / sr * 80)
    out = np.tanh(out) * 0.85
    pcm = (out * 32767).astype("<i2").tobytes()
    with wave.open(path_wav, "wb") as wf:
        wf.setnchannels(1)
        wf.setsampwidth(2)
        wf.setframerate(sr)
        wf.writeframes(pcm)


# ─────────────────────── ⑤ ffmpeg 합성 ───────────────────────
def _make_scene_clip(png: str, narration: str, dur: float, out: str):
    """정지 이미지 + (선택) 나레이션 → 줌 모션이 들어간 장면 클립."""
    frames = max(int(dur * FPS), FPS)
    zoom = f"zoompan=z='min(1+0.0009*on,1.12)':d={frames}:x='iw/2-(iw/zoom/2)':y='ih/2-(ih/zoom/2)':s={W}x{H}:fps={FPS}"
    cmd = [_ffmpeg(), "-y", "-loop", "1", "-t", f"{dur:.3f}", "-i", png]
    if narration:
        cmd += ["-i", narration,
                "-filter_complex", f"[0:v]{zoom}[v];[1:a]apad[a]",
                "-map", "[v]", "-map", "[a]", "-shortest"]
    else:
        cmd += ["-f", "lavfi", "-t", f"{dur:.3f}", "-i", "anullsrc=r=44100:cl=mono",
                "-filter_complex", f"[0:v]{zoom}[v]",
                "-map", "[v]", "-map", "1:a"]
    cmd += ["-c:v", "libx264", "-pix_fmt", "yuv420p", "-r", str(FPS),
            "-c:a", "aac", "-ar", "44100", "-ac", "1", "-t", f"{dur:.3f}", out]
    subprocess.run(cmd, check=True, capture_output=True, timeout=300)


def build_video(rows, script: dict, workdir: str, tts_voice: str) -> str:
    """장면 클립들 생성 → 이어붙이기 → BGM 믹싱. 최종 mp4 경로 반환."""
    scenes = []  # (png, narration mp3 or None, duration)

    def scene(name, render_fn, text, min_dur):
        png = os.path.join(workdir, f"{name}.png")
        render_fn(png)
        mp3 = os.path.join(workdir, f"{name}.mp3")
        dur = min_dur
        nar = None
        if text and synth_tts(text, mp3, tts_voice):
            alen = _audio_duration(mp3)
            # 나레이션이 장면 한도보다 길면: 그 문장만 말 속도를 올려 재합성
            # → 어떤 문장도 중간에 끊기지 않음 (전체 영상은 쇼츠 60초 미만 유지)
            if alen > 0 and alen + 0.45 > SCENE_MAX:
                faster = min(TTS_RATE * (alen + 0.6) / SCENE_MAX, TTS_RATE_MAX)
                print(f"  (나레이션 {alen:.1f}s > 한도 {SCENE_MAX}s → 속도 {faster:.2f}배로 재합성)")
                if synth_tts(text, mp3, tts_voice, rate=faster):
                    alen = _audio_duration(mp3)
            if alen > 0:
                dur = min(max(alen + 0.45, min_dur), SCENE_MAX)
                nar = mp3
        scenes.append((png, nar, dur))

    scene("intro", render_intro, script["intro"], INTRO_MIN)
    for i, r in enumerate(rows):
        scene(f"job{i+1}", lambda p, r=r, i=i: render_job_scene(p, r, i + 1),
              script["jobs"][i], SCENE_MIN)
    scene("outro", render_outro, script["outro"], OUTRO_MIN)

    # 장면별 클립 생성 후 concat
    clips = []
    for i, (png, nar, dur) in enumerate(scenes):
        out = os.path.join(workdir, f"clip{i}.mp4")
        _make_scene_clip(png, nar, dur, out)
        clips.append(out)
    lst = os.path.join(workdir, "list.txt")
    with open(lst, "w") as f:
        for c in clips:
            f.write(f"file '{c}'\n")
    merged = os.path.join(workdir, "merged.mp4")
    subprocess.run([_ffmpeg(), "-y", "-f", "concat", "-safe", "0", "-i", lst,
                    "-c", "copy", merged], check=True, capture_output=True, timeout=300)

    # BGM 준비 (assets/bgm.mp3 우선, 없으면 자체 합성)
    total = sum(d for _, _, d in scenes)
    if os.path.exists(ASSET_BGM):
        bgm = ASSET_BGM
    else:
        bgm = os.path.join(workdir, "bgm.wav")
        make_bgm(bgm, total + 1)

    final = os.path.join(workdir, f"올공_마감임박_TOP5_{_today_kst()}.mp4")
    subprocess.run([_ffmpeg(), "-y", "-i", merged, "-stream_loop", "-1", "-i", bgm,
                    "-filter_complex",
                    "[1:a]volume=0.16,afade=t=out:st=%.2f:d=1.2[bg];"
                    "[0:a][bg]amix=inputs=2:duration=first:dropout_transition=0[a]" % max(total - 1.2, 0),
                    "-map", "0:v", "-map", "[a]",
                    "-c:v", "copy", "-c:a", "aac", "-b:a", "128k", "-shortest", final],
                   check=True, capture_output=True, timeout=300)
    return final


# ─────────────────────── ⑥-1 Google Drive 업로드 ───────────────────────
# ※ 구글 정책상 서비스계정(로봇 계정)은 개인 '내 드라이브'에 파일을 소유할 수 없음
#   (storageQuotaExceeded). → 소유자 본인 계정의 OAuth(1회 인증, drive.file 범위)로 업로드.
#   1회 설정: Cloud Shell 에서 `python3 drive_auth.py` 실행 (README '쇼츠' 섹션 참고)
DRIVE_SCOPE = "https://www.googleapis.com/auth/drive.file"


def _drive_service_user(cfg):
    """_config/shorts 에 저장된 소유자 OAuth 자격증명으로 Drive 클라이언트 생성."""
    cid = cfg.get("drive_client_id")
    csec = cfg.get("drive_client_secret")
    rtok = cfg.get("drive_refresh_token")
    if not (cid and csec and rtok):
        return None
    from google.oauth2.credentials import Credentials
    from googleapiclient.discovery import build
    creds = Credentials(
        None, refresh_token=rtok, token_uri="https://oauth2.googleapis.com/token",
        client_id=cid, client_secret=csec, scopes=[DRIVE_SCOPE],
    )
    return build("drive", "v3", credentials=creds, cache_discovery=False)


def _ensure_app_folder(svc, ref, cfg) -> str:
    """이 앱이 만든 쇼츠 폴더를 찾거나 생성해 ID 반환 (+ Firestore 에 저장).
    drive.file 범위는 '앱이 만든 파일/폴더'만 다룰 수 있으므로, 앱 소유 폴더를 사용.
    (생성된 폴더는 드라이브에서 원하는 위치로 옮겨도 계속 사용 가능 — ID 불변)"""
    name = cfg.get("drive_folder_name") or "올공 쇼츠"
    res = svc.files().list(
        q=f"name = '{name}' and mimeType = 'application/vnd.google-apps.folder' and trashed = false",
        fields="files(id,name)", pageSize=5).execute()
    if res.get("files"):
        fid = res["files"][0]["id"]
    else:
        f = svc.files().create(body={"name": name, "mimeType": "application/vnd.google-apps.folder"},
                               fields="id").execute()
        fid = f["id"]
        print(f"  📁 드라이브에 '{name}' 폴더를 새로 만들었습니다. "
              "드라이브에서 이 폴더를 원하는 위치([공기업 브레인넷] 등)로 옮겨두시면 계속 그곳에 저장됩니다.")
    ref.set({"drive_folder_id": fid}, merge=True)
    return fid


def upload_to_drive(mp4_path: str, cfg: dict, ref):
    """소유자 OAuth 로 드라이브에 업로드. 성공 시 파일 ID 반환."""
    try:
        from googleapiclient.errors import HttpError
        from googleapiclient.http import MediaFileUpload
        svc = _drive_service_user(cfg)
        if svc is None:
            print("  ⚠️ 드라이브 개인 인증 미설정 → Cloud Shell 에서 `python3 drive_auth.py` 를 "
                  "한 번 실행해 주세요 (README '쇼츠' 섹션 참고)")
            return None

        def _create(fid):
            meta = {"name": os.path.basename(mp4_path), "parents": [fid]}
            media = MediaFileUpload(mp4_path, mimetype="video/mp4", resumable=True)
            return svc.files().create(body=meta, media_body=media, fields="id,webViewLink").execute()

        folder_id = cfg.get("drive_folder_id")
        if folder_id:
            try:
                f = _create(folder_id)
                print(f"  ✅ 드라이브 업로드 완료: {f.get('webViewLink')}")
                return f.get("id")
            except HttpError as e:
                if e.resp.status not in (403, 404):
                    raise
                # 앱이 만들지 않은 폴더(권한 없음) → 앱 폴더로 자동 전환
                print("  (지정 폴더에 권한 없음 → 앱 전용 폴더로 전환)")
        fid = _ensure_app_folder(svc, ref, cfg)
        f = _create(fid)
        print(f"  ✅ 드라이브 업로드 완료: {f.get('webViewLink')}")
        return f.get("id")
    except Exception as e:  # noqa: BLE001
        print(f"  ⚠️ 드라이브 업로드 실패: {e}")
        return None


# ─────────────────────── ⑥-1b 스크립트 Word 파일 → 드라이브 '[script]' 폴더 ───────────────────────
def build_script_docx(script: dict, rows, path: str):
    """그날의 나레이션 스크립트를 Word(.docx) 문서로 생성."""
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    doc.add_heading(f"올공 쇼츠 스크립트 — 마감임박 TOP 5 ({_today_kst()})", level=1)
    doc.add_paragraph("")
    doc.add_heading("인트로", level=2)
    doc.add_paragraph(script["intro"])
    for i, (line, r) in enumerate(zip(script["jobs"], rows), start=1):
        doc.add_heading(f"{i}위 — {r.get('기관명', '')}", level=2)
        doc.add_paragraph(line)
        meta = doc.add_paragraph(
            f"(공고: {r.get('채용공고제목', '')} / 고용유형: {r.get('고용유형', '-')}"
            f" / 마감: {_dday(r.get('공고종료일'))})")
        for run in meta.runs:
            run.font.size = Pt(9)
    doc.add_heading("아웃트로", level=2)
    doc.add_paragraph(script["outro"])
    doc.save(path)


def upload_script_docx(docx_path: str, cfg: dict, ref):
    """'[script]' 하위 폴더(없으면 쇼츠 폴더 아래 자동 생성)에 Word 파일 업로드."""
    try:
        from googleapiclient.http import MediaFileUpload
        svc = _drive_service_user(cfg)
        if svc is None:
            print("  (드라이브 개인 인증 미설정 → 스크립트 저장 건너뜀)")
            return None
        # 쇼츠 폴더 확보 → 그 아래 '[script]' 폴더 확보
        parent = cfg.get("drive_folder_id") or _ensure_app_folder(svc, ref, cfg)
        sub = cfg.get("drive_script_folder_id")
        if not sub:
            name = cfg.get("drive_script_folder_name") or "[script]"
            res = svc.files().list(
                q=(f"name = '{name}' and mimeType = 'application/vnd.google-apps.folder'"
                   f" and '{parent}' in parents and trashed = false"),
                fields="files(id)", pageSize=5).execute()
            if res.get("files"):
                sub = res["files"][0]["id"]
            else:
                f = svc.files().create(
                    body={"name": name, "mimeType": "application/vnd.google-apps.folder",
                          "parents": [parent]}, fields="id").execute()
                sub = f["id"]
            ref.set({"drive_script_folder_id": sub}, merge=True)
        meta = {"name": os.path.basename(docx_path), "parents": [sub]}
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        media = MediaFileUpload(docx_path, mimetype=mime)
        f = svc.files().create(body=meta, media_body=media, fields="id,webViewLink").execute()
        print(f"  ✅ 스크립트 문서 저장 완료: {f.get('webViewLink')}")
        return f.get("id")
    except Exception as e:  # noqa: BLE001
        print(f"  ⚠️ 스크립트 문서 저장 실패: {e}")
        return None


# ─────────────────────── ⑥-2 인스타그램 릴스 게시 ───────────────────────
def _upload_public(mp4_path: str) -> str:
    """릴스 게시용 공개 URL 확보 (Firebase Storage)."""
    import main as _crawler
    bucket = _crawler.storage.bucket()
    blob = bucket.blob(f"social/shorts-{_today_kst()}.mp4")
    blob.upload_from_filename(mp4_path, content_type="video/mp4")
    blob.make_public()
    return blob.public_url


def post_reels(db, mp4_path: str, rows) -> bool:
    """기존 _config/social 의 인스타 토큰으로 릴스 게시."""
    IG_API = "https://graph.instagram.com/v21.0"
    try:
        cfg = db.collection("_config").document("social").get()
        cfg = cfg.to_dict() if cfg.exists else {}
        tok = cfg.get("ig_access_token")
        if not tok:
            print("  (인스타 설정 없음 → 릴스 건너뜀)")
            return False
        uid = cfg.get("ig_user_id")
        if not uid:
            # social_post 와 동일: ig_user_id 미설정 시 토큰으로 자동 조회
            r0 = requests.get(f"{IG_API}/me", params={"fields": "user_id,username", "access_token": tok}, timeout=30)
            if not r0.ok:
                print(f"  ⚠️ 인스타 계정 조회 실패: {r0.status_code} {r0.text[:200]}")
                return False
            j0 = r0.json()
            uid = j0.get("user_id") or j0.get("id")
            print(f"  (인스타 계정 자동 확인: @{j0.get('username')} / {uid})")
        video_url = _upload_public(mp4_path)
        orgs = " · ".join(str(r.get("기관명", "")) for r in rows[:3])
        caption = (f"⏰ 오늘 마감임박 공공기관 채용 TOP 5!\n{orgs} 외\n"
                   f"상세 공고·지원 방법 👉 {SITE_URL}\n"
                   "#마감임박 #공공기관채용 #공기업 #채용공고 #취업 #취준 #쇼츠 #올공")
        r = requests.post(f"{IG_API}/{uid}/media", data={
            "access_token": tok, "media_type": "REELS",
            "video_url": video_url, "caption": caption, "share_to_feed": "true",
        }, timeout=60)
        r.raise_for_status()
        container = r.json().get("id")
        # 처리 완료 대기 (최대 4분)
        for _ in range(48):
            st = requests.get(f"{IG_API}/{container}",
                              params={"fields": "status_code", "access_token": tok},
                              timeout=30).json().get("status_code")
            if st == "FINISHED":
                break
            if st == "ERROR":
                print("  ⚠️ 릴스 처리 실패(ERROR)")
                return False
            time.sleep(5)
        pub = requests.post(f"{IG_API}/{uid}/media_publish",
                            data={"access_token": tok, "creation_id": container}, timeout=60)
        pub.raise_for_status()
        print(f"  ✅ 인스타 릴스 게시 완료: {pub.json().get('id')}")
        return True
    except Exception as e:  # noqa: BLE001
        print(f"  ⚠️ 릴스 게시 실패: {e}")
        return False


# ─────────────────────── 메인 진입점 ───────────────────────
def run_daily(db, df):
    """크롤링 완료 후 호출. 설정이 없으면 무동작.
    같은 날 다시 실행되면 이미 성공한 단계(드라이브/인스타)는 건너뛰고
    실패했던 단계만 재시도한다. (예: 폴더 공유를 늦게 한 경우
    스케줄러를 한 번 더 돌리면 드라이브 업로드만 다시 수행, 인스타 중복 게시 없음)"""
    ref = db.collection("_config").document("shorts")
    snap = ref.get()
    cfg = snap.to_dict() if snap.exists else {}
    if not cfg.get("enabled"):
        print("  (쇼츠 비활성 — _config/shorts.enabled=true 로 켜세요)")
        return

    today = _today_kst()
    state = cfg.get("state") if isinstance(cfg.get("state"), dict) else {}
    if state.get("date") != today:
        state = {"date": today, "drive": False, "ig": False, "script": False}
    need_drive = not state.get("drive")
    need_ig = bool(cfg.get("ig_reels", True)) and not state.get("ig")
    need_script = not state.get("script")
    if not need_drive and not need_ig and not need_script:
        print("  (오늘 쇼츠는 이미 저장·게시 완료 → 건너뜀)")
        return

    rows = select_top5(df)
    if len(rows) < 3:
        print(f"  (마감임박 공고가 {len(rows)}건뿐 → 쇼츠 생략)")
        return
    script = build_script(rows)
    tts_voice = cfg.get("tts_voice") or DEFAULT_VOICE

    with tempfile.TemporaryDirectory() as workdir:
        print("🎬 쇼츠 렌더링 시작 (마감임박 TOP 5)")
        mp4 = build_video(rows, script, workdir, tts_voice)
        size_mb = os.path.getsize(mp4) / 1e6
        print(f"  🎞️ 생성 완료: {os.path.basename(mp4)} ({size_mb:.1f}MB)")

        if need_drive:
            if upload_to_drive(mp4, cfg, ref):
                state["drive"] = True

        if need_script:
            # 그날의 나레이션 스크립트를 Word 문서로 '[script]' 폴더에 저장
            try:
                docx_path = os.path.join(workdir, f"올공_쇼츠_스크립트_{today}.docx")
                build_script_docx(script, rows, docx_path)
                if upload_script_docx(docx_path, cfg, ref):
                    state["script"] = True
            except Exception as e:  # noqa: BLE001
                print(f"  ⚠️ 스크립트 문서 생성 실패: {e}")

        if need_ig:
            if post_reels(db, mp4, rows):
                state["ig"] = True

    ref.set({"last_run": today, "state": state}, merge=True)
